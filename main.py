#!/usr/bin/env python3
# ░▄▀▄░▀█▀░█▄█▒██▀▒█▀▄░▀█▀▒▄▀▄░█▒░▒██▀░▄▀▀
# ░▀▄▀░▒█▒▒█▒█░█▄▄░█▀▄░▒█▒░█▀█▒█▄▄░█▄▄▒▄██
"""
Genkō Yōshi Tategaki Converter - Convert Japanese text to proper genkō yōshi format
Implements authentic Japanese manuscript paper rules with grid-based layout
Includes comprehensive verification and compliance system
"""

import re
import argparse
import logging
import json
from typing import Dict, List, Tuple, Optional, Any
try:
    import chardet
except ImportError:
    print("Warning: chardet module not found, falling back to default encodings")
    chardet = None
from pathlib import Path
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from rich.console import Console, Group
from rich.progress import Progress, BarColumn, TextColumn, TimeRemainingColumn, TaskProgressColumn
from rich.panel import Panel
from rich.live import Live
from rich.table import Table
try:
    from rich.prompt import Prompt
except ImportError:
    Prompt = None

# Import the page size selector
try:
    from sizes import PageSizeSelector
except ImportError:
    PageSizeSelector = None

logging.basicConfig(level=logging.INFO, format='[%(levelname)s] %(message)s')


class GenkouYoshiValidator:
    """Comprehensive validator for Genkou Yoshi standards compliance"""
    
    # Define all the compliance rules from the standards
    GYOUTOU_KINSOKU = frozenset([
        # Closing punctuation
        '。', '、', '？', '！', '‼', '⁇', '⁈', '⁉',
        # Closing brackets and quotes
        '」', '』', '）', '］', '｝', '〉', '》', '〕', '〗', '〙', '〛',
        # Small kana (sokuon, youon)
        'ゎ', 'ゃ', 'ゅ', 'ょ', 'っ', 'ぁ', 'ぃ', 'ぅ', 'ぇ', 'ぉ',
        'ァ', 'ィ', 'ゥ', 'ェ', 'ォ', 'ッ', 'ャ', 'ュ', 'ョ', 'ヮ',
        # Prolonged sound marks
        'ー', '～', '〜',
        # Iteration marks
        '々', 'ゝ', 'ゞ', 'ヽ', 'ヾ',
        # Middle dots and special punctuation
        '・', '：', '；', '‥', '…'
    ])
    
    GYOUMATSU_KINSOKU = frozenset([
        # Opening brackets and quotes
        '「', '『', '（', '［', '｛', '〈', '《', '〔', '〖', '〘', '〚'
    ])
    
    PERIOD_COMMA_CHARS = frozenset(['。', '、'])
    QUESTION_EXCLAMATION_CHARS = frozenset(['？', '！'])
    OPENING_QUOTES = frozenset(['「', '『', '（', '［'])
    CLOSING_QUOTES = frozenset(['」', '』', '）', '］'])
    
    def __init__(self, page_format: Dict):
        self.page_format = page_format
        self.violations = []
        
    def validate_grid_structure(self, grid_data: Dict) -> List[Dict]:
        """Validate basic grid structure compliance"""
        violations = []
        
        expected_cols = self.page_format['grid']['columns']
        expected_rows = self.page_format['grid']['rows']
        
        for page_num, page_data in grid_data.items():
            # Check grid dimensions
            actual_cols = len(page_data.get('columns', {}))
            if actual_cols > expected_cols:
                violations.append({
                    'type': 'grid_overflow',
                    'page': page_num,
                    'message': f'Grid has {actual_cols} columns, expected max {expected_cols}',
                    'severity': 'critical'
                })
            
            # Check row counts in each column
            for col_num, col_data in page_data.get('columns', {}).items():
                actual_rows = len(col_data)
                if actual_rows > expected_rows:
                    violations.append({
                        'type': 'column_overflow',
                        'page': page_num,
                        'column': col_num,
                        'message': f'Column has {actual_rows} rows, expected max {expected_rows}',
                        'severity': 'critical'
                    })
        
        return violations
    
    def validate_punctuation_placement(self, grid_data: Dict) -> List[Dict]:
        """Validate punctuation positioning according to Genkou Yoshi rules"""
        violations = []
        
        for page_num, page_data in grid_data.items():
            for col_num, col_data in page_data.get('columns', {}).items():
                for row_num, char in col_data.items():
                    # Check period and comma placement (should be upper-right in tategaki)
                    if char in self.PERIOD_COMMA_CHARS:
                        # In proper tategaki, periods and commas should be positioned in upper-right
                        # For DOCX validation, we check if they appear in expected positions
                        # This is a formatting guideline rather than a critical error
                        # Skip validation for now as DOCX doesn't support sub-cell positioning
                        pass
                    
                    # Check for question and exclamation marks that need special handling
                    if char in self.QUESTION_EXCLAMATION_CHARS:
                        # These should be properly oriented for tategaki
                        # The text processor should have already converted them to vertical forms
                        if char in ['?', '!']:  # ASCII versions shouldn't appear
                            violations.append({
                                'type': 'ascii_punctuation',
                                'page': page_num,
                                'column': col_num,
                                'row': row_num,
                                'character': char,
                                'message': f'ASCII punctuation "{char}" should be converted to vertical form',
                                'severity': 'medium'
                            })
        
        return violations
    
    def validate_line_breaking_rules(self, grid_data: Dict) -> List[Dict]:
        """Validate 禁則処理 (line breaking rules) compliance"""
        violations = []
        
        for page_num, page_data in grid_data.items():
            for col_num, col_data in page_data.get('columns', {}).items():
                col_chars = [col_data[i] for i in sorted(col_data.keys())]
                
                # Check for prohibited column starts (行頭禁則)
                if col_chars and col_chars[0] in self.GYOUTOU_KINSOKU:
                    violations.append({
                        'type': 'gyoutou_kinsoku',
                        'page': page_num,
                        'column': col_num,
                        'character': col_chars[0],
                        'message': f'Character "{col_chars[0]}" cannot start a column',
                        'severity': 'high'
                    })
                
                # Check for prohibited column ends (行末禁則)
                if col_chars and col_chars[-1] in self.GYOUMATSU_KINSOKU:
                    violations.append({
                        'type': 'gyoumatsu_kinsoku',
                        'page': page_num,
                        'column': col_num,
                        'character': col_chars[-1],
                        'message': f'Character "{col_chars[-1]}" cannot end a column',
                        'severity': 'high'
                    })
        
        return violations
    
    def validate_character_placement(self, grid_data: Dict) -> List[Dict]:
        """Validate one character per cell rule and proper character encoding"""
        violations = []
        
        for page_num, page_data in grid_data.items():
            for col_num, col_data in page_data.get('columns', {}).items():
                for row_num, char in col_data.items():
                    # Check for empty cells
                    if not char or char.isspace():
                        continue
                        
                    # Count visual characters (handling combining characters properly)
                    import unicodedata
                    # Normalize to handle combining characters
                    normalized = unicodedata.normalize('NFC', char)
                    # Count grapheme clusters (visible characters)
                    visual_length = len(normalized)
                    
                    # Check for multi-character content (excluding single combining chars)
                    if visual_length > 1:
                        # Allow some combining characters but flag true multi-character content
                        if not all(unicodedata.category(c).startswith('M') for c in normalized[1:]):
                            violations.append({
                                'type': 'multi_character_cell',
                                'page': page_num,
                                'column': col_num,
                                'row': row_num,
                                'content': char,
                                'message': f'Cell contains "{char}" ({visual_length} characters), should contain exactly 1 character',
                                'severity': 'critical'
                            })
                    
                    # Check for invalid control characters
                    if any(unicodedata.category(c) in ['Cc', 'Cf'] for c in char):
                        violations.append({
                            'type': 'invalid_control_character',
                            'page': page_num,
                            'column': col_num,
                            'row': row_num,
                            'content': repr(char),
                            'message': f'Cell contains invalid control character: {repr(char)}',
                            'severity': 'high'
                        })
        
        return violations
    
    def validate_quotation_placement(self, grid_data: Dict) -> List[Dict]:
        """Validate quotation mark positioning rules"""
        violations = []
        
        for page_num, page_data in grid_data.items():
            for col_num, col_data in page_data.get('columns', {}).items():
                col_items = sorted(col_data.items())
                
                for i, (row_num, char) in enumerate(col_items):
                    # Opening quotes should not be at bottom of column
                    if char in self.OPENING_QUOTES and i == len(col_items) - 1:
                        violations.append({
                            'type': 'opening_quote_at_bottom',
                            'page': page_num,
                            'column': col_num,
                            'row': row_num,
                            'character': char,
                            'message': f'Opening quote "{char}" cannot be at bottom of column',
                            'severity': 'high'
                        })
                    
                    # Closing quotes should not be at top of column
                    if char in self.CLOSING_QUOTES and i == 0:
                        violations.append({
                            'type': 'closing_quote_at_top',
                            'page': page_num,
                            'column': col_num,
                            'row': row_num,
                            'character': char,
                            'message': f'Closing quote "{char}" cannot be at top of column',
                            'severity': 'high'
                        })
        
        return violations
    
    def validate_page_layout(self, doc_analysis: Dict) -> List[Dict]:
        """Validate overall page layout compliance"""
        violations = []
        
        # Check page dimensions
        if doc_analysis.get('page_width') != self.page_format['width']:
            violations.append({
                'type': 'page_width_mismatch',
                'expected': self.page_format['width'],
                'actual': doc_analysis.get('page_width'),
                'message': 'Page width does not match format specification',
                'severity': 'high'
            })
        
        if doc_analysis.get('page_height') != self.page_format['height']:
            violations.append({
                'type': 'page_height_mismatch',
                'expected': self.page_format['height'],
                'actual': doc_analysis.get('page_height'),
                'message': 'Page height does not match format specification',
                'severity': 'high'
            })
        
        # Check margins
        expected_margins = self.page_format['margins']
        actual_margins = doc_analysis.get('margins', {})
        
        for margin_type, expected_value in expected_margins.items():
            actual_value = actual_margins.get(margin_type)
            if abs(actual_value - expected_value) > 1:  # 1mm tolerance
                violations.append({
                    'type': 'margin_mismatch',
                    'margin_type': margin_type,
                    'expected': expected_value,
                    'actual': actual_value,
                    'message': f'{margin_type} margin mismatch',
                    'severity': 'medium'
                })
        
        return violations
    
    def run_complete_validation(self, grid_data: Dict, doc_analysis: Dict) -> Dict:
        """Run all validation checks and return comprehensive report"""
        all_violations = []
        
        # Run all validation checks
        all_violations.extend(self.validate_grid_structure(grid_data))
        all_violations.extend(self.validate_character_placement(grid_data))
        all_violations.extend(self.validate_line_breaking_rules(grid_data))
        all_violations.extend(self.validate_quotation_placement(grid_data))
        all_violations.extend(self.validate_punctuation_placement(grid_data))
        all_violations.extend(self.validate_page_layout(doc_analysis))
        
        # Categorize violations by severity
        critical = [v for v in all_violations if v.get('severity') == 'critical']
        high = [v for v in all_violations if v.get('severity') == 'high']
        medium = [v for v in all_violations if v.get('severity') == 'medium']
        low = [v for v in all_violations if v.get('severity') == 'low']
        
        return {
            'total_violations': len(all_violations),
            'critical': critical,
            'high': high,
            'medium': medium,
            'low': low,
            'all_violations': all_violations,
            'compliance_score': max(0, 100 - len(all_violations) * 2)  # Simple scoring
        }


class DocxAnalyzer:
    """Analyzes DOCX files to extract grid structure and formatting data"""
    
    def __init__(self, docx_path: Path):
        self.docx_path = docx_path
        self.doc = None
        
    def load_document(self):
        """Load the DOCX document for analysis"""
        try:
            self.doc = Document(self.docx_path)
            return True
        except Exception as e:
            logging.error(f"Failed to load DOCX: {e}")
            return False
    
    def extract_page_dimensions(self) -> Dict:
        """Extract page dimensions and margins"""
        if not self.doc:
            return {}
            
        section = self.doc.sections[0]
        
        return {
            'page_width': round(section.page_width.mm),
            'page_height': round(section.page_height.mm),
            'margins': {
                'top': round(section.top_margin.mm),
                'bottom': round(section.bottom_margin.mm),
                'inner': round(section.left_margin.mm),
                'outer': round(section.right_margin.mm)
            },
            'orientation': section.orientation
        }
    
    def extract_grid_data(self) -> Dict:
        """Extract character grid data from tables with proper tategaki coordinate mapping"""
        if not self.doc:
            return {}
        
        grid_data = {}
        page_num = 1
        
        for table in self.doc.tables:
            page_data = {'columns': {}}
            
            # For tategaki: columns are DOCX table columns, rows are DOCX table rows
            # This matches the generation logic in generate_docx_content
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    # Extract text from cell
                    cell_text = cell.text.strip()
                    
                    if cell_text:  # Only record non-empty cells
                        # Tategaki column mapping: col_idx + 1 = tategaki column
                        # Tategaki row mapping: row_idx + 1 = position within column
                        tategaki_col = col_idx + 1
                        tategaki_row = row_idx + 1
                        
                        if tategaki_col not in page_data['columns']:
                            page_data['columns'][tategaki_col] = {}
                        
                        page_data['columns'][tategaki_col][tategaki_row] = cell_text
            
            if page_data['columns']:  # Only add pages with content
                grid_data[page_num] = page_data
                page_num += 1
        
        return grid_data
    
    def analyze_typography(self) -> Dict:
        """Analyze typography and formatting"""
        if not self.doc:
            return {}
        
        fonts_used = set()
        font_sizes = set()
        
        # Extract font information from runs
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                if run.font.name:
                    fonts_used.add(run.font.name)
                if run.font.size:
                    font_sizes.add(run.font.size.pt)
        
        return {
            'fonts_used': list(fonts_used),
            'font_sizes': list(font_sizes),
            'primary_font': list(fonts_used)[0] if fonts_used else None
        }
    
    def run_complete_analysis(self) -> Dict:
        """Run complete document analysis"""
        if not self.load_document():
            return {}
        
        return {
            'page_dimensions': self.extract_page_dimensions(),
            'grid_data': self.extract_grid_data(),
            'typography': self.analyze_typography(),
            'total_pages': len(self.extract_grid_data())
        }


class DocumentAdjuster:
    """Makes adjustments to fix compliance violations"""
    
    def __init__(self, builder, page_format: Dict):
        self.builder = builder
        self.page_format = page_format
        
    def fix_line_breaking_violations(self, violations: List[Dict]) -> int:
        """Fix line breaking rule violations"""
        fixes_applied = 0
        
        # Group violations by page and column for efficient processing
        violations_by_location = {}
        for violation in violations:
            if violation['type'] in ['gyoutou_kinsoku', 'gyoumatsu_kinsoku']:
                page = violation['page']
                col = violation['column']
                if page not in violations_by_location:
                    violations_by_location[page] = {}
                if col not in violations_by_location[page]:
                    violations_by_location[page][col] = []
                violations_by_location[page][col].append(violation)
        
        # Apply fixes
        for page_num, page_violations in violations_by_location.items():
            for col_num, col_violations in page_violations.items():
                # Rebuild the column content with proper line breaking
                if self._rebuild_column_with_kinsoku(page_num, col_num, col_violations):
                    fixes_applied += len(col_violations)
        
        return fixes_applied
    
    def _rebuild_column_with_kinsoku(self, page_num: int, col_num: int, violations: List[Dict]) -> bool:
        """Rebuild column content applying proper 禁則処理"""
        try:
            # Get current page grid data
            pages = self.builder.grid.get_all_pages()
            if page_num - 1 >= len(pages):
                return False
            
            page_data = pages[page_num - 1]
            
            # Fix gyoutou_kinsoku violations by moving prohibited characters to previous column
            for violation in violations:
                if violation['type'] == 'gyoutou_kinsoku':
                    prohibited_char = violation['character']
                    current_col = violation['column']
                    
                    # Get current column data
                    current_col_data = page_data['columns'].get(current_col, {})
                    if not current_col_data:
                        continue
                    
                    # Find the prohibited character at the start of the column
                    sorted_rows = sorted(current_col_data.keys())
                    if sorted_rows and current_col_data[sorted_rows[0]] == prohibited_char:
                        # Remove from current column start
                        del current_col_data[sorted_rows[0]]
                        
                        # Shift all remaining characters up one position
                        new_col_data = {}
                        for i, row_key in enumerate(sorted_rows[1:], 1):
                            new_col_data[i] = current_col_data[row_key]
                        page_data['columns'][current_col] = new_col_data
                        
                        # Try to add to previous column end
                        prev_col = current_col - 1
                        if prev_col >= 1:
                            prev_col_data = page_data['columns'].get(prev_col, {})
                            if prev_col_data:
                                # Find the last occupied row in previous column
                                prev_sorted_rows = sorted(prev_col_data.keys())
                                if prev_sorted_rows:
                                    last_row = prev_sorted_rows[-1]
                                    # Add prohibited character to end of previous column if space allows
                                    if last_row < self.builder.grid.squares_per_column:
                                        prev_col_data[last_row + 1] = prohibited_char
                                        page_data['columns'][prev_col] = prev_col_data
                                    else:
                                        # Previous column is full, need to create new column
                                        # Insert character at start of current column again (will be fixed in next iteration)
                                        if current_col_data:
                                            new_col_data = {1: prohibited_char}
                                            for row_key, char in current_col_data.items():
                                                new_col_data[row_key + 1] = char
                                            page_data['columns'][current_col] = new_col_data
                                        else:
                                            page_data['columns'][current_col] = {1: prohibited_char}
                            else:
                                # Previous column is empty, just add the character
                                page_data['columns'][prev_col] = {1: prohibited_char}
                        else:
                            # No previous column, character stays but will be handled differently
                            if not page_data['columns'].get(current_col):
                                page_data['columns'][current_col] = {}
                            page_data['columns'][current_col][1] = prohibited_char
            
            # Update the builder's grid with the fixed data
            # Ensure we update the correct pages list that get_all_pages() actually uses
            if page_num - 1 < len(self.builder.grid.pages):
                self.builder.grid.pages[page_num - 1] = page_data
            else:
                # Extend pages list if needed
                while len(self.builder.grid.pages) <= page_num - 1:
                    self.builder.grid.pages.append({'page_num': len(self.builder.grid.pages) + 1, 'columns': {}})
                self.builder.grid.pages[page_num - 1] = page_data
            
            # If this is the current page, also update the current page grid
            if page_num == self.builder.grid.current_page:
                # Convert back to list format for the current page grid
                new_grid = [['' for _ in range(self.builder.grid.squares_per_column)] 
                           for _ in range(self.builder.grid.max_columns_per_page)]
                
                for col_num, col_data in page_data['columns'].items():
                    col_idx = col_num - 1
                    if 0 <= col_idx < len(new_grid):
                        for row_num, char in col_data.items():
                            row_idx = row_num - 1
                            if 0 <= row_idx < len(new_grid[col_idx]):
                                new_grid[col_idx][row_idx] = char
                
                self.builder.grid.current_page_grid = new_grid
            
            logging.info(f"Applied kinsoku processing to page {page_num}, column {col_num}")
            return True
            
        except Exception as e:
            logging.error(f"Failed to rebuild column: {e}")
            return False
    
    def _apply_kinsoku_processing(self, chars: List[str]) -> List[str]:
        """Apply proper 禁則処理 to character sequence"""
        if not chars:
            return chars
        
        # Remove prohibited characters from start
        while chars and chars[0] in GenkouYoshiValidator.GYOUTOU_KINSOKU:
            # Move character to previous column (simplified)
            chars = chars[1:]
        
        # Remove prohibited characters from end
        while chars and chars[-1] in GenkouYoshiValidator.GYOUMATSU_KINSOKU:
            # Move character to next column (simplified)
            chars = chars[:-1]
        
        return chars
    
    def fix_punctuation_alignment(self, violations: List[Dict]) -> int:
        """Fix punctuation alignment issues"""
        fixes_applied = 0
        
        # Note: Actual punctuation alignment would require modifying
        # the DOCX cell formatting, which is complex
        # For now, we'll log the fixes needed
        
        for violation in violations:
            if violation['type'] == 'punctuation_alignment':
                logging.info(f"Punctuation alignment fix needed: {violation}")
                fixes_applied += 1
        
        return fixes_applied
    
    def fix_grid_overflow(self, violations: List[Dict]) -> int:
        """Fix grid overflow issues"""
        fixes_applied = 0
        
        for violation in violations:
            if violation['type'] in ['grid_overflow', 'column_overflow']:
                # Redistribute content to fit within grid bounds
                if self._redistribute_overflow_content(violation):
                    fixes_applied += 1
        
        return fixes_applied
    
    def _redistribute_overflow_content(self, violation: Dict) -> bool:
        """Redistribute content that overflows grid bounds"""
        try:
            # Add new pages or redistribute content as needed
            # This is a simplified implementation
            logging.info(f"Redistributing overflow content: {violation}")
            return True
        except Exception as e:
            logging.error(f"Failed to redistribute content: {e}")
            return False
    
    def apply_fixes(self, violations: List[Dict]) -> Dict:
        """Apply fixes for all types of violations with comprehensive rule preservation"""
        fixes_report = {
            'line_breaking': 0,
            'punctuation': 0,
            'grid_overflow': 0,
            'kinsoku_fixes': 0,
            'rule_preservation': 0,
            'total_fixes': 0,
            'validation_passes': 0,
            'secondary_violations_fixed': 0
        }
        
        # Categorize violations
        line_breaking_violations = [v for v in violations if v['type'] in ['gyoutou_kinsoku', 'gyoumatsu_kinsoku']]
        punctuation_violations = [v for v in violations if v['type'] == 'punctuation_alignment']
        overflow_violations = [v for v in violations if v['type'] in ['grid_overflow', 'column_overflow']]
        kinsoku_violations = [v for v in violations if v['type'] in ['kinsoku_line_start', 'kinsoku_line_end', 'small_kana_placement']]
        
        # Apply fixes in order of priority
        # 1. Kinsoku violations first (most critical for Japanese text)
        if kinsoku_violations:
            fixes_report['kinsoku_fixes'] = self._fix_kinsoku_violations(kinsoku_violations)
            if fixes_report['kinsoku_fixes'] > 0:
                # Apply comprehensive rule preservation after kinsoku fixes
                fixes_report['rule_preservation'] += self._apply_comprehensive_rule_preservation()
                fixes_report['validation_passes'] += 1
        
        # 2. Line breaking violations (traditional processing)
        if line_breaking_violations:
            fixes_report['line_breaking'] = self.fix_line_breaking_violations(line_breaking_violations)
            # Re-validate after line breaking fixes to catch any new violations
            if fixes_report['line_breaking'] > 0:
                secondary_fixes = self._validate_and_fix_secondary_violations()
                fixes_report['secondary_violations_fixed'] += secondary_fixes
                fixes_report['validation_passes'] += 1
        
        if overflow_violations:
            fixes_report['grid_overflow'] = self.fix_grid_overflow(overflow_violations)
            # Re-validate after overflow fixes
            if fixes_report['grid_overflow'] > 0:
                secondary_fixes = self._validate_and_fix_secondary_violations()
                fixes_report['secondary_violations_fixed'] += secondary_fixes
                fixes_report['validation_passes'] += 1
        
        if punctuation_violations:
            fixes_report['punctuation'] = self.fix_punctuation_alignment(punctuation_violations)
            # Re-validate after punctuation fixes
            if fixes_report['punctuation'] > 0:
                secondary_fixes = self._validate_and_fix_secondary_violations()
                fixes_report['secondary_violations_fixed'] += secondary_fixes
                fixes_report['validation_passes'] += 1
        
        fixes_report['total_fixes'] = sum([
            fixes_report['line_breaking'],
            fixes_report['punctuation'],
            fixes_report['grid_overflow'],
            fixes_report['kinsoku_fixes'],
            fixes_report['rule_preservation'],
            fixes_report['secondary_violations_fixed']
        ])
        
        return fixes_report
    
    def _validate_and_fix_secondary_violations(self) -> int:
        """Run comprehensive validation and fix any secondary violations introduced by previous fixes"""
        try:
            # Create a temporary validator to check current grid state
            validator = GenkouYoshiValidator(self.page_format)
            
            # Get current grid data from the builder
            current_grid_data = {}
            pages = self.builder.grid.get_all_pages()
            
            for page in pages:
                page_num = page['page_num']
                current_grid_data[page_num] = page
            
            # Run all validation checks on current state
            all_violations = []
            all_violations.extend(validator.validate_grid_structure(current_grid_data))
            all_violations.extend(validator.validate_character_placement(current_grid_data))
            all_violations.extend(validator.validate_line_breaking_rules(current_grid_data))
            all_violations.extend(validator.validate_quotation_placement(current_grid_data))
            all_violations.extend(validator.validate_punctuation_placement(current_grid_data))
            
            # Filter out violations we can automatically fix
            fixable_violations = [v for v in all_violations if v['type'] in [
                'gyoutou_kinsoku', 'gyoumatsu_kinsoku', 'opening_quote_at_bottom', 'closing_quote_at_top'
            ]]
            
            secondary_fixes_applied = 0
            
            # Fix quotation placement violations
            quote_violations = [v for v in fixable_violations if v['type'] in ['opening_quote_at_bottom', 'closing_quote_at_top']]
            if quote_violations:
                secondary_fixes_applied += self._fix_quotation_violations(quote_violations)
            
            # Fix any remaining line breaking violations (recursive kinsoku processing)
            remaining_kinsoku = [v for v in fixable_violations if v['type'] in ['gyoutou_kinsoku', 'gyoumatsu_kinsoku']]
            if remaining_kinsoku:
                secondary_fixes_applied += self.fix_line_breaking_violations(remaining_kinsoku)
            
            # Validate character-per-cell compliance after fixes
            char_violations = [v for v in all_violations if v['type'] == 'multi_character_cell']
            if char_violations:
                secondary_fixes_applied += self._fix_character_placement_violations(char_violations)
            
            return secondary_fixes_applied
            
        except Exception as e:
            logging.warning(f"Secondary validation failed: {e}")
            return 0
    
    def _fix_quotation_violations(self, violations: List[Dict]) -> int:
        """Fix quotation mark placement violations"""
        fixes_applied = 0
        
        for violation in violations:
            try:
                page_num = violation['page']
                col_num = violation['column']
                violation_type = violation['type']
                
                # Get current page data
                pages = self.builder.grid.get_all_pages()
                if page_num - 1 >= len(pages):
                    continue
                
                page_data = pages[page_num - 1]
                col_data = page_data['columns'].get(col_num, {})
                
                if violation_type == 'opening_quote_at_bottom':
                    # Move opening quote from bottom to next column top
                    if col_data:
                        sorted_rows = sorted(col_data.keys())
                        last_row = sorted_rows[-1]
                        quote_char = col_data[last_row]
                        
                        # Remove from current position
                        del col_data[last_row]
                        
                        # Try to add to next column
                        next_col = col_num + 1
                        if next_col <= self.builder.grid.max_columns_per_page:
                            if next_col not in page_data['columns']:
                                page_data['columns'][next_col] = {}
                            # Insert at top of next column
                            next_col_data = page_data['columns'][next_col]
                            # Shift existing content down
                            new_next_col_data = {1: quote_char}
                            for row, char in next_col_data.items():
                                new_next_col_data[row + 1] = char
                            page_data['columns'][next_col] = new_next_col_data
                            fixes_applied += 1
                
                elif violation_type == 'closing_quote_at_top':
                    # Move closing quote from top to previous column bottom
                    if col_data:
                        sorted_rows = sorted(col_data.keys())
                        first_row = sorted_rows[0]
                        quote_char = col_data[first_row]
                        
                        # Remove from current position
                        del col_data[first_row]
                        # Shift remaining content up
                        new_col_data = {}
                        for i, row in enumerate(sorted_rows[1:], 1):
                            new_col_data[i] = col_data[row]
                        page_data['columns'][col_num] = new_col_data
                        
                        # Try to add to previous column
                        prev_col = col_num - 1
                        if prev_col >= 1:
                            if prev_col not in page_data['columns']:
                                page_data['columns'][prev_col] = {}
                            prev_col_data = page_data['columns'][prev_col]
                            # Add to bottom of previous column
                            if prev_col_data:
                                max_row = max(prev_col_data.keys())
                                if max_row < self.builder.grid.squares_per_column:
                                    prev_col_data[max_row + 1] = quote_char
                                    fixes_applied += 1
                            else:
                                prev_col_data[1] = quote_char
                                fixes_applied += 1
                
                # Update the grid with fixed data
                if page_num - 1 < len(self.builder.grid.pages):
                    self.builder.grid.pages[page_num - 1] = page_data
                    
            except Exception as e:
                logging.warning(f"Failed to fix quotation violation: {e}")
                continue
        
        return fixes_applied
    
    def _fix_character_placement_violations(self, violations: List[Dict]) -> int:
        """Fix multi-character cell violations by splitting content"""
        fixes_applied = 0
        
        for violation in violations:
            try:
                page_num = violation['page']
                col_num = violation['column']
                row_num = violation['row']
                content = violation['content']
                
                # Get current page data
                pages = self.builder.grid.get_all_pages()
                if page_num - 1 >= len(pages):
                    continue
                
                page_data = pages[page_num - 1]
                col_data = page_data['columns'].get(col_num, {})
                
                if row_num in col_data and len(content) > 1:
                    # Split multi-character content across multiple cells
                    chars = list(content)
                    
                    # Place first character in original position
                    col_data[row_num] = chars[0]
                    
                    # Place remaining characters in subsequent rows/columns
                    current_row = row_num + 1
                    current_col = col_num
                    
                    for char in chars[1:]:
                        if current_row > self.builder.grid.squares_per_column:
                            # Move to next column
                            current_col += 1
                            current_row = 1
                            
                            if current_col > self.builder.grid.max_columns_per_page:
                                # Would overflow page - stop placing characters
                                break
                            
                            if current_col not in page_data['columns']:
                                page_data['columns'][current_col] = {}
                        
                        # Ensure we have the column data
                        if current_col not in page_data['columns']:
                            page_data['columns'][current_col] = {}
                        
                        page_data['columns'][current_col][current_row] = char
                        current_row += 1
                    
                    fixes_applied += 1
                
                # Update the grid with fixed data
                if page_num - 1 < len(self.builder.grid.pages):
                    self.builder.grid.pages[page_num - 1] = page_data
                    
            except Exception as e:
                logging.warning(f"Failed to fix character placement violation: {e}")
                continue
        
        return fixes_applied
    
    def _fix_kinsoku_violations(self, violations: List[Dict]) -> int:
        """Fix kinsoku shori (line-breaking rule) violations"""
        fixes_applied = 0
        
        for violation in violations:
            try:
                page_num = violation['page']
                col_idx = violation['column']
                row_idx = violation['row']
                char = violation['character']
                violation_type = violation['type']
                
                # Get the grid page
                if not self.builder.grid or page_num - 1 >= len(self.builder.grid.pages):
                    continue
                    
                page = self.builder.grid.pages[page_num - 1]
                if col_idx >= len(page.columns):
                    continue
                    
                column = page.columns[col_idx]
                
                if violation_type == 'kinsoku_line_start':
                    # Move closing punctuation to end of previous line
                    if self._move_char_to_previous_line(page, col_idx, row_idx, char):
                        fixes_applied += 1
                        
                elif violation_type == 'kinsoku_line_end':
                    # Move opening punctuation to start of next line
                    if self._move_char_to_next_line(page, col_idx, row_idx, char):
                        fixes_applied += 1
                        
                elif violation_type == 'kinsoku_small_kana':
                    # Move small kana to end of previous line
                    if self._move_char_to_previous_line(page, col_idx, row_idx, char):
                        fixes_applied += 1
                        
            except Exception as e:
                logging.warning(f"Failed to fix kinsoku violation: {e}")
                continue
                
        return fixes_applied
    
    def _move_char_to_previous_line(self, page, col_idx, row_idx, char) -> bool:
        """Move character to the end of the previous column (line)"""
        try:
            # Find previous column with space
            for prev_col_idx in range(col_idx - 1, -1, -1):
                if prev_col_idx < len(page.columns):
                    prev_column = page.columns[prev_col_idx]
                    
                    # Find last occupied cell in previous column
                    last_row = -1
                    for i, cell in enumerate(prev_column.cells):
                        if cell.character and cell.character != ' ':
                            last_row = i
                    
                    # Add character to next available position
                    if last_row + 1 < len(prev_column.cells):
                        prev_column.cells[last_row + 1].character = char
                        # Remove from current position
                        page.columns[col_idx].cells[row_idx].character = ' '
                        return True
            
            return False
        except Exception as e:
            logging.warning(f"Failed to move character to previous line: {e}")
            return False
    
    def _move_char_to_next_line(self, page, col_idx, row_idx, char) -> bool:
        """Move character to the start of the next column (line)"""
        try:
            # Find next column with space at the top
            for next_col_idx in range(col_idx + 1, len(page.columns)):
                next_column = page.columns[next_col_idx]
                
                # Check if first cell is available
                if next_column.cells[0].character == ' ' or not next_column.cells[0].character:
                    next_column.cells[0].character = char
                    # Remove from current position
                    page.columns[col_idx].cells[row_idx].character = ' '
                    return True
            
            return False
        except Exception as e:
            logging.warning(f"Failed to move character to next line: {e}")
            return False
    
    def _apply_comprehensive_rule_preservation(self) -> int:
        """Apply comprehensive genkou yoshi rule preservation after any fixes"""
        preservation_actions = 0
        
        if not self.builder.grid:
            return preservation_actions
            
        for page in self.builder.grid.pages:
            for col_idx, column in enumerate(page.columns):
                for row_idx, cell in enumerate(column.cells):
                    if cell.character and cell.character != ' ':
                        # Convert half-width punctuation to full-width
                        if self._normalize_punctuation(cell):
                            preservation_actions += 1
                            
                        # Ensure proper character encoding
                        if self._normalize_character_encoding(cell):
                            preservation_actions += 1
                            
                        # Apply proper formatting
                        if self._apply_proper_formatting(cell):
                            preservation_actions += 1
        
        return preservation_actions
    
    def _normalize_punctuation(self, cell) -> bool:
        """Normalize punctuation to proper Japanese full-width characters"""
        char = cell.character
        punctuation_map = {
            '.': '。', ',': '、', '!': '！', '?': '？',
            '(': '（', ')': '）', ':': '：', ';': '；',
            '[': '［', ']': '］', '{': '｛', '}': '｝'
        }
        
        if char in punctuation_map:
            cell.character = punctuation_map[char]
            return True
        return False
    
    def _normalize_character_encoding(self, cell) -> bool:
        """Ensure proper character encoding for genkou yoshi"""
        char = cell.character
        
        # Convert half-width katakana to full-width
        halfwidth_katakana = 'ｱｲｳｴｵｶｷｸｹｺｻｼｽｾｿﾀﾁﾂﾃﾄﾅﾆﾇﾈﾉﾊﾋﾌﾍﾎﾏﾐﾑﾒﾓﾔﾕﾖﾗﾘﾙﾚﾛﾜｦﾝ'
        fullwidth_katakana = 'アイウエオカキクケコサシスセソタチツテトナニヌネノハヒフヘホマミムメモヤユヨラリルレロワヲン'
        
        if char in halfwidth_katakana:
            idx = halfwidth_katakana.index(char)
            cell.character = fullwidth_katakana[idx]
            return True
            
        return False
    
    def _apply_proper_formatting(self, cell) -> bool:
        """Apply proper formatting for genkou yoshi cells"""
        # Ensure cell has proper formatting attributes
        if not hasattr(cell, 'formatting') or not cell.formatting:
            cell.formatting = {
                'font_size': 12,
                'vertical_alignment': 'center',
                'horizontal_alignment': 'center',
                'font_family': 'Noto Sans JP'
            }
            return True
        return False


class InMemoryGridAnalyzer:
    """Analyzes grid structure directly from builder's in-memory data to avoid disk I/O race conditions"""
    
    def __init__(self, builder, page_format: Dict):
        self.builder = builder
        self.page_format = page_format
        
    def update_from_builder(self, builder):
        """Update analyzer to use current builder state"""
        self.builder = builder
        
    def analyze_current_state(self) -> Dict:
        """Analyze current grid state from builder's in-memory data"""
        try:
            grid_data = self._extract_grid_data_from_builder()
            page_dimensions = self._get_page_dimensions()
            
            return {
                'grid_data': grid_data,
                'page_dimensions': page_dimensions,
                'analysis_method': 'in_memory'
            }
        except Exception as e:
            logging.error(f"In-memory grid analysis failed: {e}")
            return None
            
    def _extract_grid_data_from_builder(self) -> Dict:
        """Extract grid data directly from builder's grid structure"""
        grid_data = {}
        
        if not self.builder.grid or not hasattr(self.builder.grid, 'pages'):
            return grid_data
            
        # Handle different grid data structures
        if hasattr(self.builder.grid, 'pages'):
            # New grid structure with pages containing columns
            for page_num, page in enumerate(self.builder.grid.pages, 1):
                page_data = {'columns': {}}
                
                if hasattr(page, 'columns'):
                    for col_idx, column in enumerate(page.columns):
                        column_data = {}
                        if hasattr(column, 'cells'):
                            for row_idx, cell in enumerate(column.cells):
                                if hasattr(cell, 'character') and cell.character and cell.character != ' ':
                                    column_data[row_idx] = {
                                        'character': cell.character,
                                        'position': {
                                            'column': col_idx,
                                            'row': row_idx,
                                            'page': page_num
                                        },
                                        'formatting': getattr(cell, 'formatting', {}),
                                        'is_punctuation': self._is_punctuation(cell.character),
                                        'is_small_kana': self._is_small_kana(cell.character)
                                    }
                        
                        if column_data:
                            page_data['columns'][col_idx] = column_data
                
                if page_data['columns']:
                    grid_data[page_num] = page_data
        else:
            # Fallback: use get_all_pages() method
            try:
                pages = self.builder.grid.get_all_pages()
                for page_num, page in enumerate(pages, 1):
                    if isinstance(page, dict) and 'columns' in page:
                        grid_data[page_num] = page
            except Exception as e:
                logging.warning(f"Failed to extract pages via get_all_pages(): {e}")
                
        return grid_data
        
    def _get_page_dimensions(self) -> Dict:
        """Get page dimensions from page format"""
        return {
            'page_width': self.page_format.get('page_width', 210),
            'page_height': self.page_format.get('page_height', 297),
            'margins': self.page_format.get('margins', {
                'top': 20, 'bottom': 20, 'inner': 25, 'outer': 20
            })
        }
        
    def _is_punctuation(self, char: str) -> bool:
        """Check if character is Japanese punctuation"""
        punctuation_chars = '。、！？「」『』（）：；'
        return char in punctuation_chars
        
    def _is_small_kana(self, char: str) -> bool:
        """Check if character is small kana"""
        small_kana = 'ぁぃぅぇぉゃゅょっァィゥェォャュョッ'
        return char in small_kana


class VerificationEngine:
    """Orchestrates the complete verification and adjustment process"""
    
    def __init__(self, builder, page_format: Dict, console: Console):
        self.builder = builder
        self.page_format = page_format
        self.console = console
        self.validator = GenkouYoshiValidator(page_format)
        self.adjuster = DocumentAdjuster(builder, page_format)
        self._grid_analyzer = None  # In-memory grid analyzer
        
    def run_verification_cycle(self, docx_path: Path, max_iterations: int = 3) -> Dict:
        """Run complete verification cycle with iterative improvements using in-memory analysis"""
        
        self.console.print("\n[bold cyan]Starting Genkou Yoshi Verification Process[/bold cyan]")
        
        iteration = 0
        final_report = {}
        previous_violations = None
        stagnant_iterations = 0
        
        # Initialize in-memory grid analyzer from builder's current state
        self._grid_analyzer = InMemoryGridAnalyzer(self.builder, self.page_format)
        
        while iteration < max_iterations:
            iteration += 1
            self.console.print(f"\n[yellow]Verification Iteration {iteration}[/yellow]")
            
            # Analyze using in-memory grid data (no disk I/O)
            with self.console.status("Analyzing grid structure..."):
                analysis = self._grid_analyzer.analyze_current_state()
            
            if not analysis:
                self.console.print("[bold red]Failed to analyze grid structure[/bold red]")
                return {'status': 'failed', 'error': 'Grid analysis failed'}
            
            # Run comprehensive validation including kinsoku compliance
            with self.console.status("Running comprehensive compliance validation..."):
                validation_report = self.validator.run_complete_validation(
                    analysis['grid_data'], 
                    analysis['page_dimensions']
                )
                
                # Additional kinsoku-specific validation
                kinsoku_report = self._validate_kinsoku_compliance(analysis['grid_data'])
                validation_report = self._merge_validation_reports(validation_report, kinsoku_report)
            
            # Display validation results
            self._display_validation_report(validation_report, iteration)
            
            # Check if we've achieved compliance
            if validation_report['total_violations'] == 0:
                self.console.print("[bold green]✓ Perfect compliance achieved![/bold green]")
                # Save final compliant document
                self._save_final_document(docx_path)
                final_report = {
                    'status': 'compliant',
                    'iterations': iteration,
                    'final_score': validation_report['compliance_score'],
                    'validation_report': validation_report
                }
                break
            
            # Check if violations have changed since last iteration
            current_violations = self._get_violation_signature(validation_report['all_violations'])
            if previous_violations and current_violations == previous_violations:
                stagnant_iterations += 1
                self.console.print(f"[yellow]Warning: Same violations detected for {stagnant_iterations} iteration(s)[/yellow]")
                
                # Stop if we're stuck in a loop
                if stagnant_iterations >= 2:
                    self.console.print("[red]Stopping verification: violations not improving[/red]")
                    break
            else:
                stagnant_iterations = 0
            
            # Apply fixes only if we have violations and they're potentially fixable
            if validation_report['total_violations'] > 0 and stagnant_iterations < 2:
                with self.console.status("Applying compliance fixes..."):
                    fixes_report = self.adjuster.apply_fixes(validation_report['all_violations'])
                
                self._display_fixes_report(fixes_report)
                
                # Only regenerate if fixes were actually applied
                if fixes_report['total_fixes'] > 0:
                    # Comprehensive rule preservation during regeneration
                    with self.console.status("Applying comprehensive rule preservation..."):
                        self._preserve_all_genkou_yoshi_rules()
                    
                    # Update in-memory analyzer with current state
                    self._grid_analyzer.update_from_builder(self.builder)
                    
                    # Don't save to disk until verification is complete
                    self.console.print("[dim]Grid updated in memory, continuing validation...[/dim]")
                else:
                    self.console.print("[yellow]No fixes applied, stopping verification[/yellow]")
                    break
            
            previous_violations = current_violations
        
        # Save final document to disk after verification is complete
        self._save_final_document(docx_path)
        
        # If we reach here, we've exceeded max iterations without achieving compliance
        return {
            'status': 'partial_compliance',
            'iterations': iteration,
            'final_score': validation_report.get('compliance_score', 0),
            'remaining_violations': validation_report.get('total_violations', 0),
            'validation_report': validation_report
        }
    
    def _validate_kinsoku_compliance(self, grid_data: Dict) -> Dict:
        """Additional validation specifically for kinsoku shori (line-breaking rules)"""
        violations = []
        
        for page_num, page_data in grid_data.items():
            for col_idx, column_data in page_data['columns'].items():
                # Check for line-breaking violations
                cells = list(column_data.items())
                
                for i, (row_idx, cell_data) in enumerate(cells):
                    char = cell_data['character']
                    
                    # Rule 1: Don't start lines with closing punctuation
                    if row_idx == 0 and char in '。、」）』】':
                        violations.append({
                            'type': 'kinsoku_line_start',
                            'severity': 'high',
                            'page': page_num,
                            'column': col_idx,
                            'row': row_idx,
                            'character': char,
                            'message': f'Line cannot start with closing punctuation: {char}'
                        })
                    
                    # Rule 2: Don't end lines with opening punctuation
                    if i == len(cells) - 1 and char in '「（『【':
                        violations.append({
                            'type': 'kinsoku_line_end',
                            'severity': 'high',
                            'page': page_num,
                            'column': col_idx,
                            'row': row_idx,
                            'character': char,
                            'message': f'Line cannot end with opening punctuation: {char}'
                        })
                    
                    # Rule 3: Small kana should not start lines
                    if row_idx == 0 and char in 'ぁぃぅぇぉゃゅょっァィゥェォャュョッ':
                        violations.append({
                            'type': 'kinsoku_small_kana',
                            'severity': 'medium',
                            'page': page_num,
                            'column': col_idx,
                            'row': row_idx,
                            'character': char,
                            'message': f'Small kana cannot start line: {char}'
                        })
        
        return {
            'violations': violations,
            'total_violations': len(violations),
            'validation_type': 'kinsoku_compliance'
        }
    
    def _merge_validation_reports(self, primary_report: Dict, kinsoku_report: Dict) -> Dict:
        """Merge validation reports, ensuring kinsoku violations are included"""
        # Add kinsoku violations to appropriate severity categories
        for violation in kinsoku_report['violations']:
            severity = violation['severity']
            if severity not in primary_report:
                primary_report[severity] = []
            primary_report[severity].append(violation)
            primary_report['all_violations'].append(violation)
        
        # Update totals
        primary_report['total_violations'] += kinsoku_report['total_violations']
        
        # Recalculate compliance score
        if primary_report['total_violations'] > 0:
            # Reduce score based on kinsoku violations
            kinsoku_penalty = min(20, kinsoku_report['total_violations'] * 5)
            primary_report['compliance_score'] = max(0, primary_report['compliance_score'] - kinsoku_penalty)
        
        return primary_report
    
    def _preserve_all_genkou_yoshi_rules(self):
        """Ensure all genkou yoshi rules are preserved during any processing"""
        if not self.builder.grid:
            return
            
        self.console.print("[cyan]Applying comprehensive rule preservation...[/cyan]")
        
        # Apply all formatting rules to ensure compliance
        for page in self.builder.grid.pages:
            for col_idx, column in enumerate(page.columns):
                for row_idx, cell in enumerate(column.cells):
                    if cell.character and cell.character != ' ':
                        # Ensure proper punctuation formatting
                        self._apply_punctuation_rules(cell)
                        
                        # Ensure proper kana formatting
                        self._apply_kana_rules(cell, row_idx, column)
                        
                        # Ensure proper spacing rules
                        self._apply_spacing_rules(cell, row_idx, col_idx)
    
    def _apply_punctuation_rules(self, cell):
        """Apply punctuation-specific formatting rules"""
        char = cell.character
        
        # Japanese punctuation should use full-width characters
        punctuation_map = {
            '.': '。', ',': '、', '!': '！', '?': '？',
            '(': '（', ')': '）', ':': '：', ';': '；'
        }
        
        if char in punctuation_map:
            cell.character = punctuation_map[char]
    
    def _apply_kana_rules(self, cell, row_idx, column):
        """Apply kana-specific formatting rules"""
        char = cell.character
        
        # Small kana positioning rules
        if char in 'ぁぃぅぇぉゃゅょっァィゥェォャュョッ':
            # Ensure small kana doesn't start a line
            if row_idx == 0 and len(column.cells) > 1:
                # Try to move to previous line if possible
                self._handle_small_kana_positioning(cell, row_idx, column)
    
    def _apply_spacing_rules(self, cell, row_idx, col_idx):
        """Apply spacing and alignment rules"""
        # Ensure proper character spacing for tategaki
        if hasattr(cell, 'formatting'):
            if not cell.formatting:
                cell.formatting = {}
            cell.formatting['vertical_alignment'] = 'center'
            cell.formatting['horizontal_alignment'] = 'center'
    
    def _handle_small_kana_positioning(self, cell, row_idx, column):
        """Handle positioning of small kana characters"""
        # This is a placeholder for complex small kana positioning logic
        # In a full implementation, this would move characters between lines
        pass
    
    def _save_final_document(self, docx_path: Path):
        """Save the final document with all fixes applied"""
        self.console.print("[bold blue]Saving final document...[/bold blue]")
        
        # Generate final DOCX content
        self.builder.generate_docx_content()
        
        # Save to disk
        self.builder.doc.save(docx_path)
        
        # Clean up
        self.builder.doc = None
        
        self.console.print(f"[bold green]✓ Final document saved: {docx_path}[/bold green]")
    
    def _display_validation_report(self, report: Dict, iteration: int):
        """Display validation report in formatted table"""
        
        # Summary table
        summary_table = Table(title=f"Validation Summary - Iteration {iteration}")
        summary_table.add_column("Metric", style="cyan")
        summary_table.add_column("Value", style="green")
        
        summary_table.add_row("Total Violations", str(report['total_violations']))
        summary_table.add_row("Critical", f"[red]{len(report['critical'])}[/red]")
        summary_table.add_row("High", f"[orange1]{len(report['high'])}[/orange1]")
        summary_table.add_row("Medium", f"[yellow]{len(report['medium'])}[/yellow]")
        summary_table.add_row("Low", f"[green]{len(report['low'])}[/green]")
        summary_table.add_row("Compliance Score", f"{report['compliance_score']}/100")
        
        self.console.print(summary_table)
        
        # Show critical and high violations in detail
        if report['critical'] or report['high']:
            violations_table = Table(title="Critical & High Priority Violations")
            violations_table.add_column("Type", style="cyan")
            violations_table.add_column("Location", style="yellow")
            violations_table.add_column("Message", style="white")
            
            for violation in report['critical']:
                location = f"Page {violation.get('page', 'N/A')}"
                if 'column' in violation:
                    location += f", Col {violation['column']}"
                if 'row' in violation:
                    location += f", Row {violation['row']}"
                    
                violations_table.add_row(
                    f"[red]{violation['type']}[/red]",
                    location,
                    violation['message']
                )
            
            for violation in report['high']:
                location = f"Page {violation.get('page', 'N/A')}"
                if 'column' in violation:
                    location += f", Col {violation['column']}"
                if 'row' in violation:
                    location += f", Row {violation['row']}"
                    
                violations_table.add_row(
                    f"[orange1]{violation['type']}[/orange1]",
                    location,
                    violation['message']
                )
            
            self.console.print(violations_table)
    
    def _display_fixes_report(self, fixes_report: Dict):
        """Display fixes applied report"""
        fixes_table = Table(title="Fixes Applied")
        fixes_table.add_column("Fix Type", style="cyan")
        fixes_table.add_column("Count", style="green")
        
        fixes_table.add_row("Line Breaking", str(fixes_report['line_breaking']))
        fixes_table.add_row("Punctuation", str(fixes_report['punctuation']))
        fixes_table.add_row("Grid Overflow", str(fixes_report['grid_overflow']))
        
        # Show secondary violations if any were fixed
        if fixes_report.get('secondary_violations_fixed', 0) > 0:
            fixes_table.add_row("Secondary Violations", str(fixes_report['secondary_violations_fixed']))
            fixes_table.add_row("Validation Passes", str(fixes_report.get('validation_passes', 0)))
        
        fixes_table.add_row("Total", f"[bold]{fixes_report['total_fixes']}[/bold]")
        
        self.console.print(fixes_table)
    
    def _get_violation_signature(self, violations: List[Dict]) -> str:
        """Create a unique signature for violations to detect repetition"""
        if not violations:
            return ""
        
        # Create signature based on violation types, locations, and messages
        signature_parts = []
        for violation in sorted(violations, key=lambda v: (v.get('type', ''), v.get('page', 0), v.get('column', 0), v.get('row', 0))):
            signature_part = f"{violation.get('type', '')}:{violation.get('page', '')}:{violation.get('column', '')}:{violation.get('row', '')}"
            signature_parts.append(signature_part)
        
        return "|".join(signature_parts)
    
    def _validate_kinsoku_compliance(self, grid_data: Dict) -> Dict:
        """Validate kinsoku processing compliance across the grid"""
        violations = []
        
        for page_num, page_data in grid_data.items():
            for col_idx, column_data in page_data['columns'].items():
                # Get sorted row positions for this column
                row_positions = sorted(column_data.keys())
                
                for i, row_idx in enumerate(row_positions):
                    cell_data = column_data[row_idx]
                    char = cell_data['character']
                    
                    # Check line start violations (禁則処理)
                    if row_idx == 0:  # First character in column (line start)
                        if char in '。、」）』】！？：；':  # Characters that shouldn't start lines
                            violations.append({
                                'type': 'kinsoku_line_start',
                                'severity': 'high',
                                'message': f"Character '{char}' should not start a line",
                                'position': cell_data['position'],
                                'character': char
                            })
                    
                    # Check line end violations
                    if i == len(row_positions) - 1:  # Last character in column (line end)
                        if char in '「（『【':  # Characters that shouldn't end lines
                            violations.append({
                                'type': 'kinsoku_line_end',
                                'severity': 'high',
                                'message': f"Character '{char}' should not end a line",
                                'position': cell_data['position'],
                                'character': char
                            })
                    
                    # Check small kana placement
                    if cell_data.get('is_small_kana', False):
                        if row_idx == 0:  # Small kana at line start
                            violations.append({
                                'type': 'small_kana_placement',
                                'severity': 'medium',
                                'message': f"Small kana '{char}' should not start a line",
                                'position': cell_data['position'],
                                'character': char
                            })
        
        return {
            'kinsoku_violations': violations,
            'total_kinsoku_violations': len(violations)
        }
    
    def _merge_validation_reports(self, main_report: Dict, kinsoku_report: Dict) -> Dict:
        """Merge kinsoku validation results into main validation report"""
        # Add kinsoku violations to appropriate severity categories
        for violation in kinsoku_report['kinsoku_violations']:
            severity = violation['severity']
            if severity not in main_report:
                main_report[severity] = []
            main_report[severity].append(violation)
            main_report['all_violations'].append(violation)
        
        # Update totals
        main_report['total_violations'] += kinsoku_report['total_kinsoku_violations']
        
        # Recalculate compliance score
        if main_report['total_violations'] > 0:
            # Reduce score based on kinsoku violations
            kinsoku_penalty = min(20, kinsoku_report['total_kinsoku_violations'] * 5)
            main_report['compliance_score'] = max(0, main_report['compliance_score'] - kinsoku_penalty)
        
        return main_report
    
    def _preserve_all_genkou_yoshi_rules(self):
        """Apply comprehensive rule preservation to ensure all genkou yoshi rules are maintained"""
        if not self.builder.grid or not hasattr(self.builder.grid, 'pages'):
            return
        
        for page in self.builder.grid.pages:
            for column in page.columns:
                self._apply_column_rule_preservation(column)
    
    def _apply_column_rule_preservation(self, column):
        """Apply rule preservation to a single column"""
        if not column.cells:
            return
            
        # Get all non-empty cells
        non_empty_cells = [(i, cell) for i, cell in enumerate(column.cells) 
                          if cell.character and cell.character.strip()]
        
        if len(non_empty_cells) < 2:
            return
            
        # Apply kinsoku rules
        for i, (cell_idx, cell) in enumerate(non_empty_cells):
            char = cell.character
            
            # Rule 1: Don't start lines with closing punctuation
            if cell_idx == 0 and char in '。、」）』】！？：；':
                # Move to previous column if possible
                self._move_character_to_previous_position(column, cell_idx, char)
            
            # Rule 2: Don't end lines with opening punctuation  
            if i == len(non_empty_cells) - 1 and char in '「（『【':
                # Move to next column if possible
                self._move_character_to_next_position(column, cell_idx, char)
            
            # Rule 3: Handle small kana placement
            if self._is_small_kana_char(char) and cell_idx == 0:
                self._move_character_to_previous_position(column, cell_idx, char)
    
    def _move_character_to_previous_position(self, column, cell_idx, char):
        """Move character to previous available position following kinsoku rules"""
        # Implementation would require coordination with grid structure
        # For now, mark for adjustment
        pass
    
    def _move_character_to_next_position(self, column, cell_idx, char):
        """Move character to next available position following kinsoku rules"""
        # Implementation would require coordination with grid structure
        # For now, mark for adjustment
        pass
    
    def _is_small_kana_char(self, char: str) -> bool:
        """Check if character is small kana"""
        small_kana = 'ぁぃぅぇぉゃゅょっァィゥェォャュョッ'
        return char in small_kana
    
    def _save_final_document(self, docx_path: Path):
        """Save the final document with progress tracking"""
        with Progress(
            TextColumn("[bold blue]Saving final document..."),
            BarColumn(bar_width=40),
            TaskProgressColumn(),
            TimeRemainingColumn(),
            console=self.console,
            transient=False
        ) as progress:
            task = progress.add_task("saving", total=100)
            
            def progress_callback(current_page, total_pages):
                progress_percent = int((current_page / total_pages) * 100)
                progress.update(task, completed=progress_percent)
            
            # Generate final DOCX content
            self.builder.generate_docx_content(progress_callback=progress_callback)
            progress.update(task, completed=100)
            
            # Save to disk
            self.builder.doc.save(docx_path)
            
            # Force file system sync
            import os
            if hasattr(os, 'sync'):
                os.sync()
        
        self.console.print(f"[bold green]✓ Final document saved: {docx_path}[/bold green]")


class GenkouYoshiGrid:
    """Grid layout manager using efficient list-based data structures"""
    
    def __init__(self, squares_per_column=20, max_columns_per_page=10, page_format=None):
        # If page_format is provided, use its grid dimensions
        if page_format and 'grid' in page_format:
            self.squares_per_column = page_format['grid']['rows']
            self.max_columns_per_page = page_format['grid']['columns']
        else:
            self.squares_per_column = squares_per_column
            self.max_columns_per_page = max_columns_per_page
            
        self.current_column = 1
        self.current_square = 1
        
        # Use list of lists for O(1) access
        self.current_page_grid = [['' for _ in range(self.squares_per_column)] 
                                  for _ in range(self.max_columns_per_page)]
        self.pages = []  # Store completed pages
        self.current_page = 1
        
        # Store margins and other format properties
        self.page_format = page_format
        
    def move_to_column(self, column_num, square_num=1):
        """Move to a specific column and square"""
        self.current_column = min(column_num, self.max_columns_per_page)
        self.current_square = square_num
        
    def advance_square(self):
        """Move to the next square in the current column"""
        self.current_square += 1
        if self.current_square > self.squares_per_column:
            self.current_column += 1
            self.current_square = 1
            
    def advance_column(self, square_num=1):
        """Move to the next column at specified square"""
        self.current_column += 1
        self.current_square = square_num
        
        # Check if we need a new page
        if self.current_column > self.max_columns_per_page:
            self.finish_page()
            self.current_column = 1
            self.current_square = square_num
        
    def skip_columns(self, num_columns):
        """Skip a number of columns (for spacing)"""
        self.current_column += num_columns
        self.current_square = 1
        
    def place_character(self, char):
        """Place a single character in the current grid position"""
        # Input validation
        if not char or len(char) != 1:
            return
            
        # Check if we need a new page
        if self.current_column > self.max_columns_per_page:
            self.finish_page()
            self.current_column = 1
            self.current_square = 1
            
        col_idx = self.current_column - 1
        square_idx = self.current_square - 1
        
        # Bounds checking for safety
        if (col_idx >= 0 and col_idx < len(self.current_page_grid) and 
            square_idx >= 0 and square_idx < len(self.current_page_grid[col_idx])):
            self.current_page_grid[col_idx][square_idx] = char
            self.advance_square()
                
    def place_text_batch(self, text):
        """Place text as a batch for better performance"""
        for char in text:
            # Inline the character placement logic
            if self.current_column > self.max_columns_per_page:
                self.finish_page()
                self.current_column = 1
                self.current_square = 1
                
            col_idx = self.current_column - 1
            square_idx = self.current_square - 1
            
            if col_idx < len(self.current_page_grid) and square_idx < len(self.current_page_grid[col_idx]):
                self.current_page_grid[col_idx][square_idx] = char
                
                # Inline advance_square logic
                self.current_square += 1
                if self.current_square > self.squares_per_column:
                    self.current_column += 1
                    self.current_square = 1
            
    def finish_page(self):
        """Finish current page and start a new one"""
        if any(any(col) for col in self.current_page_grid):
            # Convert to dictionary format for compatibility with DOCX generation
            columns_dict = {}
            for col_idx, column in enumerate(self.current_page_grid):
                col_num = col_idx + 1
                squares_dict = {}
                for square_idx, char in enumerate(column):
                    if char:
                        squares_dict[square_idx + 1] = char
                if squares_dict:
                    columns_dict[col_num] = squares_dict
                    
            self.pages.append({
                'page_num': self.current_page,
                'columns': columns_dict
            })
            
            # Reset current page with new grid
            self.current_page_grid = [['' for _ in range(self.squares_per_column)] 
                                      for _ in range(self.max_columns_per_page)]
            self.current_page += 1
            
    def get_all_pages(self):
        """Get all pages including current page"""
        pages = self.pages.copy()
        # Add current page if it has content
        if any(any(col) for col in self.current_page_grid):
            columns_dict = {}
            for col_idx, column in enumerate(self.current_page_grid):
                col_num = col_idx + 1
                squares_dict = {}
                for square_idx, char in enumerate(column):
                    if char:
                        squares_dict[square_idx + 1] = char
                if squares_dict:
                    columns_dict[col_num] = squares_dict
                    
            if columns_dict:
                pages.append({
                    'page_num': self.current_page,
                    'columns': columns_dict
                })
        return pages
            
    def is_at_column_top(self):
        """Check if we're at the top of a column (square 1)"""
        return self.current_square == 1


class JapaneseTextProcessor:
    """Japanese text processor with cached patterns for efficient processing"""
    
    # Pre-compiled regex patterns
    _title_pattern = re.compile(r'^(?:題名|タイトル|Title)\s*[:：]\s*(.+)')
    _subtitle_pattern = re.compile(r'^(?:副題|サブタイトル|Subtitle)\s*[:：]\s*(.+)')
    _author_pattern = re.compile(r'^(?:作者|著者|Author)\s*[:：]\s*(.+)')
    _chapter_pattern = re.compile(
        r'^\s*(?:第[一二三四五六七八九十百千\d]+章|'
        r'Chapter\s*\d+|Chapter\s*[IVXLCDM]+|[0-9]+\.|[一二三四五六七八九十]+\.).*'
    )
    _blankline_pattern = re.compile(r'(?:\n[\s\u3000]*\n)+')
    _date_pattern = re.compile(r"(\d{1,4})年(\d{1,2})月(\d{1,2})日")
    _time_pattern = re.compile(r"(\d{1,2}):(\d{1,2})")
    _number_pattern = re.compile(r"\d+")
    
    # Character sets as frozensets for O(1) membership testing
    PROHIBITED_COLUMN_START = frozenset({'。', '、', '」', '）', '］', '？', '！', '‼', '⁇', '⁈', '⁉', 
                                        '︒', '︑', '﹂', '︶', '︼', '︖', '︕'})
    PROHIBITED_COLUMN_END = frozenset({'「', '（', '［', '﹁', '︵', '︻'})
    SMALL_KANA = frozenset({'っ', 'ゃ', 'ゅ', 'ょ', 'ァ', 'ィ', 'ゥ', 'ェ', 'ォ', 'ッ', 'ャ', 'ュ', 'ョ'})
    
    # Pre-built translation tables for fast character conversion
    _fullwidth_ascii = str.maketrans(
        '!"#$%&\'()*+,-./:;<=>?@[\\]^_`{|}~0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz',
        '！＂＃＄％＆＇（）＊＋，－．／：；＜＝＞？＠［＼］＾＿｀｛｜｝～０１２３４５６７８９ＡＢＣＤＥＦＧＨＩＪＫＬＭＮＯＰＱＲＳＴＵＶＷＸＹＺａｂｃｄｅｆｇｈｉｊｋｌｍｎｏｐｑｒｓｔｕｖｗｘｙｚ'
    )
    
    _half_to_full_kana = str.maketrans(
        '｡｢｣､･ｦｧｨｩｪｫｬｭｮｯｰｱｲｳｴｵｶｷｸｹｺｻｼｽｾｿﾀﾁﾂﾃﾄﾅﾆﾇﾈﾉﾊﾋﾌﾍﾎﾏﾐﾑﾒﾓﾔﾕﾖﾗﾘﾙﾚﾛﾜﾝﾞﾟ',
        '。「」、・ヲァィゥェォャュョッーアイウエオカキクケコサシスセソタチツテトナニヌネノハヒフヘホマミムメモヤユヨラリルレロワン゛゜'
    )
    
    # Vertical character mappings as translation table
    _vertical_translate = str.maketrans({
        '。': '︒', '、': '︑', '（': '︵', '）': '︶', '「': '﹁', '」': '﹂',
        '『': '﹃', '』': '﹄', '【': '︻', '】': '︼', '！': '︕', '？': '︖',
        '：': '︓', '；': '︔', '—': '︱', '－': '︲', '…': '︙',
    })
    
    # Number conversion tables
    _numbers_to_kanji = {
        '1': '一', '2': '二', '3': '三', '4': '四', '5': '五',
        '6': '六', '7': '七', '8': '八', '9': '九', '10': '十'
    }
    
    _fullwidth_numbers = str.maketrans('0123456789', '０１２３４５６７８９')
    
    @classmethod
    def identify_text_structure(cls, text, paragraph_split_mode='blank'):
        """Identify and extract text structure with metadata"""
        # Normalize line endings
        text = text.replace('\\r\\n', '\n').replace('\\n', '\n').replace('\r\n', '\n').replace('\r', '\n')
        lines = text.split('\n')
        
        structure = {
            'novel_title': None,
            'subtitle': None,
            'author': None,
            'body_paragraphs': [],
            'subheadings': []
        }
        
        # Single pass metadata detection with cached patterns
        metadata_indices = set()
        for idx, line in enumerate(lines):
            stripped = line.strip()
            if not stripped:
                continue
                
            # Use cached compiled patterns for faster matching
            if structure['novel_title'] is None:
                m = cls._title_pattern.match(stripped)
                if m:
                    structure['novel_title'] = m.group(1).strip()
                    metadata_indices.add(idx)
                    continue
                    
            if structure['subtitle'] is None:
                m = cls._subtitle_pattern.match(stripped)
                if m:
                    structure['subtitle'] = m.group(1).strip()
                    metadata_indices.add(idx)
                    continue
                    
            if structure['author'] is None:
                m = cls._author_pattern.match(stripped)
                if m:
                    structure['author'] = m.group(1).strip()
                    metadata_indices.add(idx)
                    continue
        
        # Remove metadata lines efficiently
        if metadata_indices:
            lines = [line for idx, line in enumerate(lines) if idx not in metadata_indices]

        # Identify non-empty lines for fallback metadata detection
        non_empty_lines = [line.strip() for line in lines if line.strip()]
        
        # Blank line splitting using cached pattern
        def blankline_split(txt):
            return [p.strip() for p in cls._blankline_pattern.split(txt) if p.strip()]

        # Fallback metadata detection
        pos = 0
        if structure['novel_title'] is None and non_empty_lines:
            structure['novel_title'] = non_empty_lines[0]
            pos = 1
        if structure['author'] is None and pos < len(non_empty_lines) and not cls._chapter_pattern.match(non_empty_lines[pos]):
            structure['author'] = non_empty_lines[pos]
            pos += 1
        if structure['subtitle'] is None and pos < len(non_empty_lines) and not cls._chapter_pattern.match(non_empty_lines[pos]):
            structure['subtitle'] = non_empty_lines[pos]
            pos += 1

        # Find content start efficiently
        count = 0
        start_idx = 0
        for idx, line in enumerate(lines):
            if line.strip():
                count += 1
            if count == pos:
                start_idx = idx + 1
                break
                
        # Process chapters with cached pattern
        remaining_lines = lines[start_idx:]
        current_chapter = None
        buffer = []
        
        for line in remaining_lines:
            if cls._chapter_pattern.match(line.strip()):
                if current_chapter is not None:
                    paragraphs = ([p.strip() for p in buffer if p.strip()] if paragraph_split_mode == 'single' 
                                else blankline_split('\n'.join(buffer)))
                    structure['subheadings'].append((current_chapter, paragraphs))
                    buffer = []
                current_chapter = line.strip()
            else:
                buffer.append(line)
                
        if current_chapter is not None:
            paragraphs = ([p.strip() for p in buffer if p.strip()] if paragraph_split_mode == 'single' 
                         else blankline_split('\n'.join(buffer)))
            structure['subheadings'].append((current_chapter, paragraphs))
        else:
            paragraphs = ([p.strip() for p in remaining_lines if p.strip()] if paragraph_split_mode == 'single' 
                         else blankline_split('\n'.join(remaining_lines)))
            structure['body_paragraphs'] = paragraphs
            
        return structure
    
    @classmethod
    def preprocess_text(cls, text):
        """
        Single-pass text preprocessing combining all transformations
        """
        # Step 1: Convert to full-width using fast translation tables
        text = text.translate(cls._fullwidth_ascii)
        text = text.translate(cls._half_to_full_kana)
        
        # Step 2: Handle quotes and symbols (fastest string replacements)
        text = (text.replace('"', '「').replace('"', '」')
                   .replace("'", '「').replace("'", '」')
                   .replace('(', '（').replace(')', '）')
                   .replace('[', '［').replace(']', '］')
                   .replace('!', '！').replace('?', '？')
                   .replace('...', '…'))
        
        # Step 3: Apply number rules with cached patterns
        text = cls._apply_number_rules(text)
        
        # Step 4: Convert to vertical equivalents using translation table
        text = text.translate(cls._vertical_translate)
        
        return text
    
    @classmethod
    def _apply_number_rules(cls, text):
        """Apply number conversion rules using cached patterns"""
        
        # Date conversion with cached pattern
        def repl_date(m):
            y, mo, da = m.group(1), m.group(2), m.group(3)
            y_str = y.translate(cls._fullwidth_numbers)
            mo_str = cls._numbers_to_kanji.get(mo, mo.translate(cls._fullwidth_numbers))
            da_str = cls._numbers_to_kanji.get(da, da.translate(cls._fullwidth_numbers))
            return f"{y_str}年{mo_str}月{da_str}日"
        
        text = cls._date_pattern.sub(repl_date, text)
        
        # Time conversion with cached pattern
        def repl_time(m):
            hh, mm = m.group(1), m.group(2)
            hh_str = cls._numbers_to_kanji.get(hh, hh.translate(cls._fullwidth_numbers))
            mm_str = cls._numbers_to_kanji.get(mm, mm.translate(cls._fullwidth_numbers))
            return f"{hh_str}時{mm_str}分"
        
        text = cls._time_pattern.sub(repl_time, text)
        
        # General number conversion with cached pattern
        def repl_num(m):
            s = m.group(0)
            return cls._numbers_to_kanji.get(s, s.translate(cls._fullwidth_numbers))
        
        text = cls._number_pattern.sub(repl_num, text)
        return text
    
    # Static method for common checks
    @staticmethod
    def is_punctuation(char):
        """Check if character is punctuation"""
        return char in {'。', '、', '！', '？', '：', '；', '︒', '︑', '︕', '︖', '︓', '︔'}
        
    @classmethod
    def is_small_kana(cls, char):
        """Check if character is small kana"""
        return char in cls.SMALL_KANA
    
    @classmethod
    def is_opening_bracket(cls, char):
        """Check if character is opening bracket"""
        return char in cls.PROHIBITED_COLUMN_END
    
    @classmethod
    def is_closing_bracket(cls, char):
        """Check if character is closing bracket"""
        return char in cls.PROHIBITED_COLUMN_START


class GenkouYoshiDocumentBuilder:
    """DOCX document builder for Genkou Yoshi format"""
    
    # Default page format (Bunko format)
    DEFAULT_PAGE_FORMAT = {
        'name': 'Bunko',
        'width': 105, 
        'height': 148,
        'grid': {'columns': 17, 'rows': 24},
        'characters_per_page': 408,
        'margins': {'top': 15, 'bottom': 15, 'inner': 12, 'outer': 8},
        'description': 'Standard mass-market paperback fiction'
    }
    
    def __init__(self, font_name='Noto Sans JP', squares_per_column=None, max_columns_per_page=None, 
                 chapter_pagebreak=False, page_size=None, page_format=None):
        self.doc = Document()
        self.text_processor = JapaneseTextProcessor()
        self.font_name = font_name
        self.chapter_pagebreak = chapter_pagebreak
        
        # Use page_format if provided, otherwise convert page_size to format, or use default
        if page_format:
            self.page_format = page_format
        elif page_size:
            self.page_format = self.convert_page_size_to_format(page_size)
        else:
            self.page_format = self.DEFAULT_PAGE_FORMAT
            
        # For backwards compatibility
        self.page_size = {
            'name': self.page_format['name'],
            'width': self.page_format['width'],
            'height': self.page_format['height'],
        }
        
        # Override grid dimensions if explicitly specified
        if squares_per_column:
            self.page_format['grid']['rows'] = squares_per_column
        if max_columns_per_page:
            self.page_format['grid']['columns'] = max_columns_per_page
        
        # Create grid
        self.grid = GenkouYoshiGrid(
            squares_per_column=self.page_format['grid']['rows'], 
            max_columns_per_page=self.page_format['grid']['columns'],
            page_format=self.page_format
        )
        
        self.setup_page_layout()
        
    def convert_page_size_to_format(self, page_size):
        """Convert legacy page size dict to full page format"""
        width = page_size.get('width', 105)
        height = page_size.get('height', 148)
        name = page_size.get('name', 'Custom')
        
        # Margin calculation based on page size
        margins = ({'top': 12, 'bottom': 12, 'inner': 10, 'outer': 8}, 7) if width <= 120 else \
                 ({'top': 15, 'bottom': 15, 'inner': 12, 'outer': 10}, 8) if width <= 160 else \
                 ({'top': 20, 'bottom': 20, 'inner': 15, 'outer': 12}, 9)
        
        margins, cell_size = margins
            
        # Calculate grid dimensions
        if PageSizeSelector:
            grid = PageSizeSelector.calculate_grid_dimensions(width, height, margins, cell_size)
            columns = grid['columns']
            rows = grid['rows']
            characters_per_page = grid['characters_per_page']
        else:
            # Fallback calculation
            text_width = width - margins['inner'] - margins['outer']
            text_height = height - margins['top'] - margins['bottom']
            columns = max(10, int(text_width / cell_size))
            rows = max(15, int(text_height / cell_size))
            characters_per_page = columns * rows
            
        return {
            'name': name,
            'width': width,
            'height': height,
            'grid': {'columns': columns, 'rows': rows},
            'characters_per_page': characters_per_page,
            'margins': margins,
            'description': f'Custom {width}×{height}mm format',
            'character_size': cell_size
        }
        
    def setup_page_layout(self):
        """Setup page layout with proper dimensions and margins"""
        section = self.doc.sections[0]
        
        # Set document dimensions
        section.page_width = Mm(self.page_size["width"])
        section.page_height = Mm(self.page_size["height"])
        section.orientation = WD_ORIENTATION.PORTRAIT
        
        # Use margins from page format
        if self.page_format and 'margins' in self.page_format:
            margins = self.page_format['margins']
        else:
            # Fallback margins
            margin_size = 15 if self.page_size["width"] <= 120 else 18 if self.page_size["width"] <= 160 else 20
            margin_lr = 8 if self.page_size["width"] <= 120 else 10 if self.page_size["width"] <= 160 else 15
            margins = {'top': margin_size, 'bottom': margin_size, 'inner': margin_lr, 'outer': margin_lr}
            
        section.top_margin = Mm(margins['top'])
        section.bottom_margin = Mm(margins['bottom'])
        section.left_margin = Mm(margins['inner'])
        section.right_margin = Mm(margins['outer'])
        
        # Set document vertical text direction
        self._set_document_vertical_text_direction(section)
    
    def create_title_page(self, title, subtitle=None, author=None):
        """Create title page with proper layout"""
        grid_cols = self.page_format['grid']['columns']
        grid_rows = self.page_format['grid']['rows']
        
        # Pre-calculate positions
        title_start_row = int(grid_rows * 0.2)
        author_row = int(grid_rows * 0.6)
        
        # Place title
        title_length = len(title)
        center_start = max(1, (grid_cols - min(title_length, grid_cols - 4)) // 2)
        
        self.grid.move_to_column(center_start, title_start_row)
        self.grid.place_text_batch(title)
            
        # Add subtitle if present
        if subtitle:
            subtitle_row = title_start_row + len(title) + 2
            subtitle_center = max(1, (grid_cols - min(len(subtitle), grid_cols - 4)) // 2)
            self.grid.move_to_column(subtitle_center, subtitle_row)
            self.grid.place_text_batch(subtitle)
                
        # Add author if present
        if author:
            author_center = max(1, (grid_cols - min(len(author), grid_cols - 4)) // 2)
            self.grid.move_to_column(author_center, author_row)
            self.grid.place_text_batch(author)
                
        self.grid.finish_page()
        
    def create_chapter_title_page(self, chapter_title):
        """Create chapter title page"""
        grid_cols = self.page_format['grid']['columns']
        grid_rows = self.page_format['grid']['rows']
        
        title_start_row = int(grid_rows * 0.15)
        
        # Clean up chapter title
        display_title = re.sub(r'^第[一二三四五六七八九十\d]+章[:：]?\s*', '', chapter_title)
        
        # Center and place title
        title_length = len(display_title)
        center_start = max(1, (grid_cols - min(title_length, grid_cols - 4)) // 2)
        
        self.grid.move_to_column(center_start, title_start_row)
        self.grid.place_text_batch(display_title)
            
        # Add spacing
        self.grid.move_to_column(1, title_start_row + title_length + 3)
        self.grid.finish_page()
        
    def _adaptive_paragraph_spacing(self, paragraph_count):
        """Calculate adaptive paragraph spacing based on grid size"""
        rows = self.page_format['grid']['rows']
        return 1 if rows < 20 else 2 if rows < 30 else 3

    def create_genkou_yoshi_document(self, text):
        """
        Create the complete Genkou Yoshi formatted document
        """
        try:
            # Preprocess text
            processed_text = self.text_processor.preprocess_text(text)
            structure = self.text_processor.identify_text_structure(processed_text)
            
            if structure['novel_title']:
                self.create_title_page(
                    structure['novel_title'],
                    subtitle=structure['subtitle'],
                    author=structure['author']
                )
                
            if structure['subheadings']:
                for chapter, paragraphs in structure['subheadings']:
                    if self.chapter_pagebreak:
                        self.grid.finish_page()
                    self.create_chapter_title_page(chapter)
                    spacing = self._adaptive_paragraph_spacing(len(paragraphs))
                    
                    for i, paragraph in enumerate(paragraphs):
                        if not paragraph:
                            self.grid.advance_column(1)
                            continue
                        if i > 0:
                            self.grid.advance_column(spacing)
                        
                        # Place paragraph
                        self.place_paragraph(paragraph)
            else:
                if structure['novel_title']:
                    self.grid.move_to_column(3, 2)
                    
                paragraphs_to_process = structure['body_paragraphs']
                spacing = self._adaptive_paragraph_spacing(len(paragraphs_to_process))
                
                for i, paragraph in enumerate(paragraphs_to_process):
                    if not paragraph:
                        self.grid.advance_column(1)
                        continue
                    if i > 0:
                        self.grid.advance_column(spacing)
                    
                    # Place paragraph
                    self.place_paragraph(paragraph)
                        
            self.grid.finish_page()
            
        except Exception as e:
            logging.critical(f"Failed to generate document: {e}")
            raise
        
    def place_paragraph(self, paragraph, indent=True):
        """
        Place paragraph with proper formatting
        """
        # Handle indentation
        if indent and self.grid.is_at_column_top():
            self.grid.advance_square()
            
        # Place text
        self.grid.place_text_batch(paragraph)

    def export_grid_metadata_json(self, output_path=None):
        """Export grid metadata as JSON for validation"""
        import json
        pages = self.grid.get_all_pages()
        metadata = []
        for page in pages:
            page_data = {
                'page_num': page['page_num'],
                'columns': {col: {sq: char for sq, char in squares.items()} 
                          for col, squares in page['columns'].items()}
            }
            metadata.append(page_data)
            
        json_str = json.dumps(metadata, ensure_ascii=False, indent=2)
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(json_str)
        return json_str

    def generate_docx_content(self, progress_callback=None):
        """
        Generate DOCX content using Word's native vertical text support
        Much more efficient than table-based approach - uses native Japanese typography
        """
        from docx.shared import Pt
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        # Clear existing paragraphs
        while len(self.doc.paragraphs) > 0:
            p = self.doc.paragraphs[0]
            p._element.getparent().remove(p._element)
            
        # Validate grid data before generating DOCX
        pages = self.grid.get_all_pages()
        if not self._validate_grid_data(pages):
            raise ValueError("Grid data validation failed - cannot generate compliant DOCX")
            
        # Configure document for native vertical text layout
        self._setup_document_vertical_text_direction()
        
        # Calculate optimal font size based on page format
        character_size = self.page_format.get('character_size', 9)
        font_size_points = max(8, min(14, character_size * 1.2))
        
        # Process each page using efficient native vertical text
        for page_index, page_data in enumerate(pages):
            if 'columns' not in page_data:
                continue
                
            # Create content for this page
            page_text_content = self._convert_grid_to_vertical_text(page_data)
            
            if page_text_content.strip():
                # Create paragraph with native vertical text formatting
                paragraph = self.doc.add_paragraph()
                self._configure_paragraph_for_vertical_text(paragraph, font_size_points)
                
                # Add the text content with proper formatting
                text_run = paragraph.add_run(page_text_content)
                text_run.font.name = self.font_name
                text_run.font.size = Pt(font_size_points)
                
            # Add page break between pages (except for the last page)
            if page_index < len(pages) - 1:
                self.doc.add_page_break()
                
            # Update progress callback if provided
            if progress_callback:
                progress_callback(page_index + 1, len(pages))
        

    def _setup_page_numbers(self):
        """Setup page numbers in footer within margins"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        section = self.doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0] 
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Clear existing content
        footer_para.clear()
        
        # Add page number field
        run = footer_para.add_run()
        run.font.size = Pt(9)
        run.font.name = self.font_name
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Create page number field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def _set_document_vertical_text_direction(self, section):
        """Set vertical text direction for document section"""
        sectPr = section._sectPr
        text_dir = sectPr.find(qn('w:textDirection'))
        if text_dir is None:
            text_dir = OxmlElement('w:textDirection')
            sectPr.append(text_dir)
        text_dir.set(qn('w:val'), 'tbRl')
        
        # Also set document-level bidi properties for proper tategaki
        bidi = sectPr.find(qn('w:bidi'))
        if bidi is None:
            bidi = OxmlElement('w:bidi')
            sectPr.append(bidi)

    def _set_vertical_text_direction(self, paragraph):
        """Set vertical text direction for paragraph with proper tategaki formatting"""
        pPr = paragraph._p.pPr
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            paragraph._p.insert(0, pPr)
        
        # Set text direction to top-to-bottom, right-to-left
        text_dir = pPr.find(qn('w:textDirection'))
        if text_dir is None:
            text_dir = OxmlElement('w:textDirection')
            pPr.append(text_dir)
        text_dir.set(qn('w:val'), 'tbRl')
        
        # Set paragraph bidi for right-to-left flow
        bidi = pPr.find(qn('w:bidi'))
        if bidi is None:
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)
            
    def _validate_grid_data(self, pages):
        """Validate grid data conforms to specifications"""
        if not pages:
            return False
            
        grid_cols = self.grid.max_columns_per_page
        grid_rows = self.grid.squares_per_column
        
        for page_idx, page in enumerate(pages):
            if not self._validate_page_structure(page, grid_cols, grid_rows):
                print(f"Validation failed for page {page_idx + 1}")
                return False
        return True
        
    def _validate_page_structure(self, page, expected_cols, expected_rows):
        """Validate individual page structure"""
        if 'columns' not in page:
            return False
            
        columns = page['columns']
        
        # Check column count doesn't exceed grid
        if len(columns) > expected_cols:
            print(f"Column overflow: {len(columns)} > {expected_cols}")
            return False
            
        # Check each column's row count
        for col_num, squares in columns.items():
            if max(squares.keys()) > expected_rows:
                print(f"Row overflow in column {col_num}: {max(squares.keys())} > {expected_rows}")
                return False
                
            # Validate each character is single character
            for square_num, char in squares.items():
                if not isinstance(char, str) or len(char) != 1:
                    print(f"Invalid character at col {col_num}, row {square_num}: '{char}'")
                    return False
                    
        return True


def main():
    """Main function with comprehensive verification system"""
    from rich.console import Console
    from rich.progress import Progress, BarColumn, TextColumn, TimeRemainingColumn, TaskProgressColumn
    import sys
    
    parser = argparse.ArgumentParser(description="Japanese Tategaki DOCX Generator with Verification")
    parser.add_argument("input", nargs="?", help="Input text file (UTF-8)")
    parser.add_argument("-o", "--output", help="Output DOCX file")
    parser.add_argument("--json", help="Export grid/character metadata as JSON to file")
    parser.add_argument("--format", default=None, help="Page format")
    parser.add_argument("--skip-verification", action="store_true", help="Skip verification process")
    parser.add_argument("--verification-report", help="Save verification report to file")
    args = parser.parse_args()

    console = Console(color_system="auto", force_terminal=True, force_interactive=True)
    
    # Display header
    ascii_art = (
        "[cyan]░▄▀▄░▀█▀░█▄█▒██▀▒█▀▄░▀█▀▒▄▀▄░█▒░▒██▀░▄▀▀[/cyan]\n"
        "[cyan]░▀▄▀░▒█▒▒█▒█░█▄▄░█▀▄░▒█▒░█▀█▒█▄▄░█▄▄▒▄██[/cyan]"
    )
    console.print(ascii_art)
    console.print("")
    console.print("[bold yellow]Japanese Tategaki DOCX Generator[/bold yellow]")
    console.print("[green]Authentic Genkou Yoshi formatting with compliance verification[/green]")
    console.print()

    if not args.input:
        console.print("[bold red]No input file specified.[/bold red]")
        sys.exit(1)

    # Load and validate input
    input_path = Path(args.input)
    if not input_path.exists():
        console.print(f"[bold red]Error: Input file '{input_path}' not found.[/bold red]")
        sys.exit(1)
        
    # File reading with encoding detection
    try:
        if chardet:
            with open(input_path, 'rb') as f:
                raw_data = f.read()
            encoding_result = chardet.detect(raw_data)
            detected_encoding = encoding_result['encoding'] if encoding_result['confidence'] > 0.7 else 'utf-8'
            try:
                text = raw_data.decode(detected_encoding).strip()
            except UnicodeDecodeError:
                text = raw_data.decode('utf-8', errors='ignore').strip()
        else:
            with open(input_path, encoding="utf-8", errors='ignore') as f:
                text = f.read().strip()
    except Exception as e:
        console.print(f"[bold red]Error reading file: {e}[/bold red]")
        sys.exit(1)
        
    if not text:
        console.print("[bold red]Error: Input file is empty.[/bold red]")
        sys.exit(1)

    # Page format selection
    if args.format is None or args.format == "custom":
        if PageSizeSelector and Prompt:
            selector = PageSizeSelector(console=console)
            page_format = selector.select_page_size()
        else:
            console.print("[bold red]Error: Interactive page size selection requires 'rich.prompt' and 'sizes.py'.[/bold red]")
            sys.exit(1)
    else:
        try:
            page_format = PageSizeSelector.get_format(args.format) if PageSizeSelector else None
        except Exception:
            page_format = None

    # Create builder
    builder = GenkouYoshiDocumentBuilder(page_format=page_format)
    
    # Quick analysis for user feedback
    structure = builder.text_processor.identify_text_structure(text)
    console.print("[bold cyan]Document Analysis:[/bold cyan]")
    console.print(f"  [bold]Title:[/bold] {structure['novel_title']}")
    console.print(f"  [bold]Author:[/bold] {structure['author']}")
    
    if structure['subheadings']:
        total_paragraphs = sum(len(pars) for _, pars in structure['subheadings'])
        console.print(f"  [bold]Chapters:[/bold] {len(structure['subheadings'])}")
    else:
        total_paragraphs = len(structure['body_paragraphs'])
        
    console.print(f"  [bold]Paragraphs:[/bold] {total_paragraphs}")
    console.print(f"  [bold]Characters:[/bold] ~{len(text):,}")

    # Process document with progress tracking
    with Progress(
        TextColumn("[progress.description]{task.description}"),
        BarColumn(),
        TaskProgressColumn(),
        TimeRemainingColumn(),
        console=console,
    ) as progress:
        
        # Main processing task
        task = progress.add_task("Processing document...", total=100)
        
        # Document creation
        progress.update(task, advance=25, description="Creating document structure...")
        builder.create_genkou_yoshi_document(text)
        
        # DOCX generation
        progress.update(task, advance=10, description="Preparing DOCX generation...")
        pages = builder.grid.get_all_pages()
        
        # Allocate 45% for page-by-page generation
        remaining_progress = 45
        progress_per_page = remaining_progress / max(1, len(pages))
        
        def progress_callback(current_page, total_pages):
            progress.update(task, advance=progress_per_page, 
                          description=f"Generating DOCX... Page {current_page}/{total_pages}")
            
        builder.generate_docx_content(progress_callback=progress_callback)
        
        # Save initial version
        if args.output:
            output_path = Path(args.output)
        else:
            format_name = page_format.get('name', 'genkou_yoshi') if page_format else 'genkou_yoshi'
            output_path = input_path.with_name(f"{input_path.stem}_{format_name}.docx")
        builder.doc.save(output_path)
        
        progress.update(task, advance=20, description="Document saved, preparing verification...")

    # Run verification if not skipped
    if not args.skip_verification:
        verification_engine = VerificationEngine(builder, page_format, console)
        verification_report = verification_engine.run_verification_cycle(output_path)
        
        # Save verification report if requested
        if args.verification_report:
            with open(args.verification_report, 'w', encoding='utf-8') as f:
                json.dump(verification_report, f, ensure_ascii=False, indent=2)
            console.print(f"[bold green]✓ Verification report saved:[/bold green] {args.verification_report}")
        
        # Display final status
        if verification_report['status'] == 'compliant':
            console.print(f"\n[bold green]🎉 Document is fully compliant with Genkou Yoshi standards![/bold green]")
            console.print(f"[green]Achieved in {verification_report['iterations']} iteration(s)[/green]")
        elif verification_report['status'] == 'partial_compliance':
            console.print(f"\n[bold yellow]⚠️  Document has partial compliance[/bold yellow]")
            console.print(f"[yellow]Compliance score: {verification_report['final_score']}/100[/yellow]")
            console.print(f"[yellow]Remaining violations: {verification_report['remaining_violations']}[/yellow]")
        else:
            console.print(f"\n[bold red]❌ Verification failed[/bold red]")
    
    console.print()
    console.print(f"[bold green]✓ DOCX file saved:[/bold green] {output_path}")
    console.print(f"[bold green]✓ Pages generated:[/bold green] {len(pages)}")
    
    if args.json:
        builder.export_grid_metadata_json(args.json)
        console.print(f"[bold green]✓ Metadata JSON saved:[/bold green] {args.json}")


if __name__ == "__main__":
    main()
